import os
import docx
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Corporate Colors
BRAND_BLUE_RGB = RGBColor(2, 132, 199)
SECONDARY_GRAY_RGB = RGBColor(59, 81, 102)
MUTED_SLATE_RGB = RGBColor(107, 124, 147)
BG_LIGHT_HEX = "F0F9FF"                    # Light blue shading for tables
HEADER_BLUE_HEX = "0284C7"                 # Header row color
BORDER_GRAY_HEX = "CBD5E1"                 # Thin borders

def set_cell_background(cell, color_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def apply_text_styling(run, font_name="Arial", size_pt=12, bold=False, italic=False, color_rgb=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb
    else:
        run.font.color.rgb = RGBColor(0, 0, 0)

def add_heading_us(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.line_spacing = 1.15
    
    run = p.add_run(text)
    if level == 1:
        apply_text_styling(run, font_name="Arial", size_pt=14, bold=True)
    elif level == 2:
        apply_text_styling(run, font_name="Arial", size_pt=12, bold=True)
    elif level == 3:
        apply_text_styling(run, font_name="Arial", size_pt=11, bold=True, italic=True)
    return p

def add_body_paragraph(doc, text="", bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    if text:
        run = p.add_run(text)
        apply_text_styling(run, font_name="Arial", size_pt=11, bold=bold, italic=italic)
    return p

def add_instruction_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    apply_text_styling(run, font_name="Arial", size_pt=9.5, italic=True, color_rgb=MUTED_SLATE_RGB)
    return p

def create_table_borders(table):
    tblPr = table._tbl.tblPr
    borders_xml = f'''
    <w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="{BORDER_GRAY_HEX}"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{BORDER_GRAY_HEX}"/>
        <w:left w:val="single" w:sz="4" w:space="0" w:color="{BORDER_GRAY_HEX}"/>
        <w:right w:val="single" w:sz="4" w:space="0" w:color="{BORDER_GRAY_HEX}"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{BORDER_GRAY_HEX}"/>
        <w:insideV w:val="none"/>
    </w:tblBorders>
    '''
    tblPr.append(parse_xml(borders_xml))

def main():
    doc = docx.Document()
    
    # Configure US Letter margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # --- TITLE PAGE ---
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_p.paragraph_format.space_before = Pt(24)
    logo_p.paragraph_format.space_after = Pt(24)
    if os.path.exists("Logo Moinhos.png"):
        logo_p.add_run().add_picture("Logo Moinhos.png", width=Inches(1.8))
        
    inst_p = doc.add_paragraph()
    inst_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inst_p.paragraph_format.space_after = Pt(72)
    inst_p.paragraph_format.line_spacing = 1.15
    run_inst1 = inst_p.add_run("US HEALTHCARE INNOVATION LAB\n")
    apply_text_styling(run_inst1, font_name="Arial", size_pt=12, bold=True)
    run_inst2 = inst_p.add_run("MBA IN ARTIFICIAL INTELLIGENCE APPLIED TO HEALTH\nCOURSE: RISK ASSESSMENT & RETURN ON INVESTMENT (ROI)")
    apply_text_styling(run_inst2, font_name="Arial", size_pt=11)
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(12)
    run_title = title_p.add_run("BUSINESS CASE, RISK AND RETURN (ROI) REPORT")
    apply_text_styling(run_title, font_name="Arial", size_pt=16, bold=True)
    
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(72)
    run_sub = sub_p.add_run("Executive Summary & Final AI Investment Proposal")
    apply_text_styling(run_sub, font_name="Arial", size_pt=12, italic=True, color_rgb=SECONDARY_GRAY_RGB)
    
    # Metadata Block
    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    
    meta_labels = [
        ("Team / Group:", "Group _____"),
        ("Members:", "1. ____________________\n2. ____________________\n3. ____________________\n4. ____________________\n5. ____________________")
    ]
    for idx, (label, val) in enumerate(meta_labels):
        row = meta_table.rows[idx]
        cell_lbl = row.cells[0]
        cell_val = row.cells[1]
        
        p_lbl = cell_lbl.paragraphs[0]
        run_lbl = p_lbl.add_run(label)
        apply_text_styling(run_lbl, font_name="Arial", size_pt=11, bold=True)
        
        p_val = cell_val.paragraphs[0]
        run_val = p_val.add_run(val)
        apply_text_styling(run_val, font_name="Arial", size_pt=11)
        
        cell_lbl.width = Inches(1.8)
        cell_val.width = Inches(4.7)
        
    doc.add_paragraph().paragraph_format.space_before = Pt(80)
    
    city_p = doc.add_paragraph()
    city_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    city_p.paragraph_format.space_after = Pt(0)
    run_city = city_p.add_run("UNITED STATES\n2026")
    apply_text_styling(run_city, font_name="Arial", size_pt=12, bold=True)
    
    doc.add_page_break()
    
    # --- 1. EXECUTIVE SUMMARY ---
    add_heading_us(doc, "1. Executive Summary", level=1)
    add_instruction_paragraph(doc, 
        "Instructions: Write a concise, high-impact summary of up to 200 words. "
        "Briefly outline the identified clinical/operational bottleneck, the proposed AI-based solution, "
        "the total capital required, the financial simulation results (NPV, Payback), and the final investment recommendation."
    )
    add_body_paragraph(doc, "[Write your team's Executive Summary here...]")
    
    # --- 2. CLINICAL SETTING & OPERATIONAL BOTTLENECK ---
    add_heading_us(doc, "2. Clinical Setting & Operational Bottleneck", level=1)
    add_instruction_paragraph(doc, 
        "Instructions: Based on your team's assigned case scenario, describe the clinical or operational bottleneck "
        "present in the baseline workflow. Detail the workflow inefficiencies (e.g., triage wait times, manual pharmacy dispensing errors) "
        "and the associated financial, regulatory, or clinical liability consequences before implementation."
    )
    add_body_paragraph(doc, "[Describe the institution's baseline workflow, volumes, and identified inefficiencies...]")
    
    # --- 3. PROPOSED ARTIFICIAL INTELLIGENCE SOLUTION ---
    add_heading_us(doc, "3. Proposed Artificial Intelligence Solution", level=1)
    add_instruction_paragraph(doc, 
        "Instructions: Present the proposed AI solution. Detail the technological scope, the data flow "
        "(how the AI integrates with the EHR/EMR or hospital hardware), and the planned clinician validation and adoption timeline."
    )
    add_body_paragraph(doc, "[Detail the solution architecture, EMR integration, and clinical workflow design...]")
    
    # --- 4. RISK CLASSIFICATION & MITIGATION MATRIX ---
    add_heading_us(doc, "4. Risk Classification & Mitigation Matrix", level=1)
    add_instruction_paragraph(doc, 
        "Instructions: Transcribe the risk mapping completed during the Class 1 dynamic. "
        "Fill in the matrix below detailing the impact of each risk dimension, the clinical/technological mitigation strategy, "
        "and the estimated annual contingency cost."
    )
    
    # Risk Table
    risk_table = doc.add_table(rows=5, cols=4)
    risk_table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    create_table_borders(risk_table)
    
    headers_risk = ["Dimension", "Mapped Risk (Impact)", "Mitigation Strategy", "Contingency Cost"]
    col_widths_risk = [Inches(1.2), Inches(2.2), Inches(2.2), Inches(1.2)]
    
    for idx, text in enumerate(headers_risk):
        cell = risk_table.rows[0].cells[idx]
        cell.width = col_widths_risk[idx]
        p = cell.paragraphs[0]
        run = p.add_run(text)
        apply_text_styling(run, font_name="Arial", size_pt=10, bold=True, color_rgb=RGBColor(255, 255, 255))
        set_cell_background(cell, HEADER_BLUE_HEX)
        
    risk_rows = [
        ("Technical", "[Integration complexity, accuracy degradation, EHR latency]", "[Model calibration, API redundancies, data backups]", "$ _________"),
        ("Operational", "[Alert fatigue, clinician workflow rejection, physical downtime]", "[Onsite training, super-user programs, backup staff]", "$ _________"),
        ("Clinical-Cultural", "[False negatives, physician distrust, ethical/bias concerns]", "[Dual validation protocols, ethics board review, audits]", "$ _________"),
        ("Financial", "[Cost overruns, maintenance escalations, claim denials]", "[Fixed-price SLAs, budget reserves, claims tracking]", "$ _________")
    ]
    
    for row_idx, data in enumerate(risk_rows):
        row = risk_table.rows[row_idx + 1]
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            cell.width = col_widths_risk[col_idx]
            p = cell.paragraphs[0]
            run = p.add_run(text)
            apply_text_styling(run, font_name="Arial", size_pt=9.5, color_rgb=SECONDARY_GRAY_RGB)
            if row_idx % 2 == 1:
                set_cell_background(cell, BG_LIGHT_HEX)
                
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # --- 5. FINANCIAL SIMULATION & STRESS TESTING ---
    add_heading_us(doc, "5. Financial Simulation & Stress Testing", level=1)
    add_instruction_paragraph(doc, 
        "Instructions: Transcribe the ROI metrics from the Class 2 dynamic. "
        "Present the results under three adoption levels (Expected, Realistic, Pessimistic) and clearly indicate the project's Break Point."
    )
    
    # Financial Inputs Block
    inputs_p = doc.add_paragraph()
    inputs_p.paragraph_format.line_spacing = 1.15
    inputs_run = inputs_p.add_run("Financial Modeling Parameters Used:\n")
    apply_text_styling(inputs_run, font_name="Arial", size_pt=11, bold=True)
    
    inputs_text = (
        "• Initial Capital Investment (CAPEX): $ __________________\n"
        "• Expected Annual Benefit (OPEX Savings/Revenue): $ __________________\n"
        "• Annual Maintenance Cost (OPEX): $ __________________\n"
        "• Annual Cost of Mapped Risks (Risk Cost): $ __________________\n"
        "• Time Horizon: _____ years"
    )
    inputs_run_val = inputs_p.add_run(inputs_text)
    apply_text_styling(inputs_run_val, font_name="Arial", size_pt=11)
    inputs_p.paragraph_format.space_after = Pt(12)
    
    # Simulation Table
    sim_table = doc.add_table(rows=5, cols=4)
    sim_table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    create_table_borders(sim_table)
    
    headers_sim = ["Financial Metric", "Expected Scenario", "Realistic Scenario", "Pessimistic Scenario"]
    col_widths_sim = [Inches(2.2), Inches(1.5), Inches(1.5), Inches(1.6)]
    
    for idx, text in enumerate(headers_sim):
        cell = sim_table.rows[0].cells[idx]
        cell.width = col_widths_sim[idx]
        p = cell.paragraphs[0]
        run = p.add_run(text)
        apply_text_styling(run, font_name="Arial", size_pt=10, bold=True, color_rgb=RGBColor(255, 255, 255))
        set_cell_background(cell, HEADER_BLUE_HEX)
        
    sim_rows = [
        ("Clinician Adoption Rate (%)", "90%", "70%", "45%"),
        ("Return on Investment (ROI %)", "________ %", "________ %", "________ %"),
        ("Payback Period (Years)", "________ years", "________ years", "________ years"),
        ("Net Present Value (NPV)", "$ ____________", "$ ____________", "$ ____________")
    ]
    
    for row_idx, data in enumerate(sim_rows):
        row = sim_table.rows[row_idx + 1]
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            cell.width = col_widths_sim[col_idx]
            p = cell.paragraphs[0]
            run = p.add_run(text)
            
            is_bold = col_idx > 0 and row_idx > 0
            color = BRAND_BLUE_RGB if is_bold else RGBColor(0,0,0)
            apply_text_styling(run, font_name="Arial", size_pt=9.5, bold=is_bold, color_rgb=color)
            
            if row_idx % 2 == 1:
                set_cell_background(cell, BG_LIGHT_HEX)
                
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    bp_p = doc.add_paragraph()
    bp_p.paragraph_format.line_spacing = 1.15
    bp_run = bp_p.add_run("Identified Break Point: ")
    apply_text_styling(bp_run, font_name="Arial", size_pt=11, bold=True)
    bp_val = bp_p.add_run("[Pessimistic Scenario / Realistic Scenario / Viable in all scenarios]\n")
    apply_text_styling(bp_val, font_name="Arial", size_pt=11)
    
    just_p = doc.add_paragraph()
    just_p.paragraph_format.line_spacing = 1.15
    just_run = just_p.add_run("Stress Sensitivity Analysis:\n")
    apply_text_styling(just_run, font_name="Arial", size_pt=11, bold=True)
    just_val = just_p.add_run("[Provide a qualitative analysis here detailing how clinician adoption rate and unmitigated risk costs degrade the Net Present Value and extend the Payback period...]")
    apply_text_styling(just_val, font_name="Arial", size_pt=11, italic=True, color_rgb=MUTED_SLATE_RGB)
    
    # --- 6. FINAL RECOMMENDATION & GOVERNANCE ---
    add_heading_us(doc, "6. Final Investment Recommendation (GO / NO-GO)", level=1)
    add_instruction_paragraph(doc, 
        "Instructions: State the team's formal recommendation to C-level executives (GO to proceed with the contract, "
        "or NO-GO to defer/reject). Justify your decision by referencing the risk-stressed ROI, clinician safety tolerances, "
        "and governance criteria."
    )
    
    dec_p = doc.add_paragraph()
    dec_p.paragraph_format.line_spacing = 1.15
    dec_run = dec_p.add_run("Formal Recommendation: ")
    apply_text_styling(dec_run, font_name="Arial", size_pt=11, bold=True)
    dec_val = dec_p.add_run("[   ] GO (APPROVED)    /    [   ] NO-GO (REJECTED)")
    apply_text_styling(dec_val, font_name="Arial", size_pt=11, bold=True, color_rgb=BRAND_BLUE_RGB)
    
    just_p2 = doc.add_paragraph()
    just_p2.paragraph_format.line_spacing = 1.15
    just_lbl = just_p2.add_run("Governance Justification:\n")
    apply_text_styling(just_lbl, font_name="Arial", size_pt=11, bold=True)
    just_val2 = just_p2.add_run(
        "[Summarize the final business case recommendation. Explain how mitigations lower risk to acceptable levels "
        "or why high integration friction and liability exposure support a NO-GO decision...]"
    )
    apply_text_styling(just_val2, font_name="Arial", size_pt=11, italic=True, color_rgb=MUTED_SLATE_RGB)

    # --- 7. REFERENCES ---
    add_heading_us(doc, "7. Bibliography & References", level=1)
    add_instruction_paragraph(doc,
        "Instructions: List the references used to support this business case. Below are core curriculum references for guidance:"
    )
    
    ref_list = [
        "ARBACHE, Fernando. Healthcare Innovation & Artificial Intelligence Applied to Risk Management. New York: Academic Press, 2024.",
        "CLEVERLEY, William O.; CLEVERLEY, James O.; SONG, Paula H. Essentials of Health Care Finance. 9th ed. Burlington: Jones & Bartlett Learning, 2023.",
        "DAMODARAN, Aswath. Investment Valuation: tools and techniques for determining the value of any asset. 4th ed. Hoboken: Wiley, 2024.",
        "HUBBARD, Douglas W. How to Measure Anything: finding the value of intangibles in business. 4th ed. Hoboken: Wiley, 2024.",
        "TOPOL, Eric. Deep Medicine: how artificial intelligence can make healthcare human again. New York: Basic Books, 2019."
    ]
    
    for ref_str in ref_list:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(ref_str)
        apply_text_styling(run, font_name="Arial", size_pt=11)

    # Save the template to the public directory
    output_dir = "/Users/faila/Library/CloudStorage/GoogleDrive-failapereira@gmail.com/My Drive/MOINHOS/us/public"
    os.makedirs(output_dir, exist_ok=True)
    file_path = os.path.join(output_dir, "Final_Project_Template.docx")
    doc.save(file_path)
    print(f"US Word Template successfully generated at: {file_path}")

if __name__ == "__main__":
    main()
