import os
import docx
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls

# Colors (kept for table headers and accents, text body is pure black ABNT NBR style)
BRAND_BLUE_RGB = RGBColor(2, 132, 199)
SECONDARY_GRAY_RGB = RGBColor(59, 81, 102)
MUTED_SLATE_RGB = RGBColor(107, 124, 147)
BG_LIGHT_HEX = "F0F9FF"                    # Light blue shading for tables
HEADER_BLUE_HEX = "0284C7"                 # Header row color
BORDER_GRAY_HEX = "CBD5E1"                 # Thin borders

def set_cell_background(cell, color_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def apply_text_styling(run, font_name="Arial", size_pt=12, bold=False, italic=False, color_rgb=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb
    else:
        run.font.color.rgb = RGBColor(0, 0, 0) # Default ABNT black

def add_heading_abnt(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.line_spacing = 1.5
    
    # ABNT headings are usually left-aligned, bold, same font
    run = p.add_run(text)
    if level == 1:
        apply_text_styling(run, font_name="Arial", size_pt=14, bold=True)
    elif level == 2:
        apply_text_styling(run, font_name="Arial", size_pt=12, bold=True)
    elif level == 3:
        apply_text_styling(run, font_name="Arial", size_pt=12, bold=True, italic=True)
    return p

def add_body_paragraph(doc, text="", bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    
    if text:
        run = p.add_run(text)
        apply_text_styling(run, font_name="Arial", size_pt=12, bold=bold, italic=italic)
    return p

def add_instruction_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Instructions do not have first line indent, they are distinct
    run = p.add_run(text)
    apply_text_styling(run, font_name="Arial", size_pt=9.5, italic=True, color_rgb=MUTED_SLATE_RGB)
    return p

def add_list_item_abnt(doc, text, is_numbered=False):
    style = 'List Number' if is_numbered else 'List Bullet'
    p = doc.add_paragraph(style=style)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    # List items do not have first line indent
    run = p.add_run(text)
    apply_text_styling(run, font_name="Arial", size_pt=12)
    return p

def create_table_borders(table):
    tblPr = table._tbl.tblPr
    borders_xml = f'''
    <w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="{BORDER_GRAY_HEX}"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{BORDER_GRAY_HEX}"/>
        <w:left w:val="single" w:sz="4" w:space="0" w:color="{BORDER_GRAY_HEX}"/>
        <w:right w:val="single" w:sz="4" w:space="0" w:color="{BORDER_GRAY_HEX}"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{BORDER_GRAY_HEX}"/>
        <w:insideV w:val="none"/>
    </w:tblBorders>
    '''
    tblPr.append(parse_xml(borders_xml))

def main():
    doc = docx.Document()
    
    # Configure ABNT margins
    # Superior: 3,0 cm | Esquerda: 3,0 cm | Inferior: 2,0 cm | Direita: 2,0 cm
    for section in doc.sections:
        section.top_margin = Cm(3.0)
        section.left_margin = Cm(3.0)
        section.bottom_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        
    # --- CAPA (ABNT) ---
    # Logo institutional
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_p.paragraph_format.space_before = Pt(0)
    logo_p.paragraph_format.space_after = Pt(12)
    if os.path.exists("Logo Moinhos.png"):
        logo_p.add_run().add_picture("Logo Moinhos.png", width=Inches(1.8))
        
    inst_p = doc.add_paragraph()
    inst_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inst_p.paragraph_format.space_after = Pt(72)
    inst_p.paragraph_format.line_spacing = 1.15
    run_inst1 = inst_p.add_run("FACULDADE DE CIÊNCIAS DA SAÚDE MOINHOS DE VENTO\n")
    apply_text_styling(run_inst1, font_name="Arial", size_pt=12, bold=True)
    run_inst2 = inst_p.add_run("MBA EM INTELIGÊNCIA ARTIFICIAL APLICADA À SAÚDE\nDISCIPLINA DE AVALIAÇÃO DE RISCOS E INVESTIMENTOS ROI")
    apply_text_styling(run_inst2, font_name="Arial", size_pt=11)
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(12)
    run_title = title_p.add_run("BUSINESS CASE E RELATÓRIO DE RISCOS E RETORNO (ROI)")
    apply_text_styling(run_title, font_name="Arial", size_pt=16, bold=True)
    
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(72)
    run_sub = sub_p.add_run("Relatório Executivo e Proposta de Investimento Final em IA na Saúde")
    apply_text_styling(run_sub, font_name="Arial", size_pt=12, italic=True, color_rgb=SECONDARY_GRAY_RGB)
    
    # Metadata Block
    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    
    meta_labels = [
        ("Equipe / Grupo:", "Grupo _____"),
        ("Integrantes:", "1. ____________________\n2. ____________________\n3. ____________________\n4. ____________________\n5. ____________________")
    ]
    for idx, (label, val) in enumerate(meta_labels):
        row = meta_table.rows[idx]
        cell_lbl = row.cells[0]
        cell_val = row.cells[1]
        
        p_lbl = cell_lbl.paragraphs[0]
        run_lbl = p_lbl.add_run(label)
        apply_text_styling(run_lbl, font_name="Arial", size_pt=11, bold=True)
        
        p_val = cell_val.paragraphs[0]
        run_val = p_val.add_run(val)
        apply_text_styling(run_val, font_name="Arial", size_pt=11)
        
        cell_lbl.width = Inches(1.8)
        cell_val.width = Inches(4.7)
        
    # Spacing to place city and year at the bottom
    doc.add_paragraph().paragraph_format.space_before = Pt(80)
    
    city_p = doc.add_paragraph()
    city_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    city_p.paragraph_format.space_after = Pt(0)
    run_city = city_p.add_run("PORTO ALEGRE\n2026")
    apply_text_styling(run_city, font_name="Arial", size_pt=12, bold=True)
    
    doc.add_page_break()
    
    # --- 1. SUMÁRIO EXECUTIVO ---
    add_heading_abnt(doc, "1. Sumário Executivo", level=1)
    add_instruction_paragraph(doc, 
        "Instruções: Redija um sumário conciso e de alto impacto de até 200 palavras. "
        "Apresente de forma sintética o problema assistencial identificado, a solução baseada em IA, "
        "o investimento total exigido, os resultados da simulação financeira e a recomendação final de investimento."
    )
    add_body_paragraph(doc, "[Escreva aqui o Sumário Executivo de sua equipe...]")
    
    # --- 2. CONTEXTO HOSPITALAR E PROBLEMA ---
    add_heading_abnt(doc, "2. Contexto Hospitalar e Gargalo Identificado", level=1)
    add_instruction_paragraph(doc, 
        "Instruções: Com base na descrição do cenário recebido pela sua equipe, descreva o gargalo clínico ou operacional "
        "existente no baseline. Detalhe as perdas assistenciais (ex: tempo excessivo de triagem, erros manuais de dispensação) "
        "e as consequências financeiras/jurídicas associadas a essa ineficiência antes do projeto."
    )
    add_body_paragraph(doc, "[Descreva o cenário baseline da instituição, volumetria e as ineficiências identificadas...]")
    
    # --- 3. SOLUÇÃO TECNOLÓGICA PROPOSTA ---
    add_heading_abnt(doc, "3. Solução Proposta em Inteligência Artificial", level=1)
    add_instruction_paragraph(doc, 
        "Instruções: Apresente a solução de IA proposta. Descreva o escopo tecnológico, o fluxo de dados "
        "(como a IA se integra aos prontuários ou hardware do hospital) e o cronograma planejado de validação e adoção "
        "pela equipe assistencial."
    )
    add_body_paragraph(doc, "[Detalhe a arquitetura da solução, integração de dados e fluxo operacional assistencial desenhado...]")
    
    # --- 4. MAPEAMENTO E MITIGAÇÃO DE RISCOS ---
    add_heading_abnt(doc, "4. Matriz de Riscos e Mitigações", level=1)
    add_instruction_paragraph(doc, 
        "Instruções: Transcreva o mapeamento detalhado efetuado na Dinâmica da Aula 1. "
        "Preencha a matriz abaixo detalhando o impacto de cada risco e o respectivo plano de contingenciamento "
        "assistencial ou tecnológico com o custo estimado."
    )
    
    # Risk Table
    risk_table = doc.add_table(rows=5, cols=4)
    risk_table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    create_table_borders(risk_table)
    
    headers_risk = ["Categoria", "Risco Mapeado (Impacto)", "Estratégia de Mitigação", "Custo Contingência"]
    col_widths_risk = [Cm(2.5), Cm(5.0), Cm(5.0), Cm(2.5)]
    
    # Style Header Row
    for idx, text in enumerate(headers_risk):
        cell = risk_table.rows[0].cells[idx]
        cell.width = col_widths_risk[idx]
        p = cell.paragraphs[0]
        run = p.add_run(text)
        apply_text_styling(run, font_name="Arial", size_pt=10, bold=True, color_rgb=RGBColor(255, 255, 255))
        set_cell_background(cell, HEADER_BLUE_HEX)
        
    risk_rows = [
        ("Técnico", "[Risco de integração, acurácia ou falha de leitura]", "[Backup de dados, APIs redundantes, calibragem]", "R$ _________"),
        ("Operacional", "[Fadiga de alarmes, rejeição, pane física]", "[Suporte técnico, treinamento de contingência]", "R$ _________"),
        ("Clínico-Cultural", "[Falsos negativos, resistência médica, ética]", "[Protocolo duplo, comitê de ética, auditoria]", "R$ _________"),
        ("Financeiro", "[Estouro de custos, inflação de insumos]", "[Contratos de longo prazo, travas orçamentárias]", "R$ _________")
    ]
    
    for row_idx, data in enumerate(risk_rows):
        row = risk_table.rows[row_idx + 1]
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            cell.width = col_widths_risk[col_idx]
            p = cell.paragraphs[0]
            run = p.add_run(text)
            apply_text_styling(run, font_name="Arial", size_pt=9.5, color_rgb=SECONDARY_GRAY_RGB)
            if row_idx % 2 == 1:
                set_cell_background(cell, BG_LIGHT_HEX)
                
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # --- 5. AVALIAÇÃO FINANCEIRA E TESTE DE ESTRESSE ---
    add_heading_abnt(doc, "5. Simulação Financeira e Testes de Estresse", level=1)
    add_instruction_paragraph(doc, 
        "Instruções: Transcreva as métricas da simulação de ROI efetuada na Dinâmica da Aula 2. "
        "Apresente os resultados sob estresse de adoção esperada, realista e pessimista. "
        "Indique com clareza o Ponto de Quebra financeira da iniciativa."
    )
    
    # Financial Inputs Block
    inputs_p = doc.add_paragraph()
    inputs_p.paragraph_format.line_spacing = 1.5
    inputs_run = inputs_p.add_run("Parâmetros do Modelo Utilizado:\n")
    apply_text_styling(inputs_run, font_name="Arial", size_pt=11, bold=True)
    
    inputs_text = (
        "• Investimento Inicial: R$ __________________\n"
        "• Benefício Anual Esperado: R$ __________________\n"
        "• Custo Anual de Manutenção: R$ __________________\n"
        "• Custo Anual de Riscos Mapeados (Matriz): R$ __________________\n"
        "• Horizonte de Análise: _____ anos"
    )
    inputs_run_val = inputs_p.add_run(inputs_text)
    apply_text_styling(inputs_run_val, font_name="Arial", size_pt=11)
    inputs_p.paragraph_format.space_after = Pt(12)
    
    # Simulation Table
    sim_table = doc.add_table(rows=5, cols=4)
    sim_table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    create_table_borders(sim_table)
    
    headers_sim = ["Indicador", "Cenário Esperado", "Cenário Realista", "Cenário Pessimista"]
    col_widths_sim = [Cm(5.0), Cm(3.3), Cm(3.3), Cm(3.4)]
    
    for idx, text in enumerate(headers_sim):
        cell = sim_table.rows[0].cells[idx]
        cell.width = col_widths_sim[idx]
        p = cell.paragraphs[0]
        run = p.add_run(text)
        apply_text_styling(run, font_name="Arial", size_pt=10, bold=True, color_rgb=RGBColor(255, 255, 255))
        set_cell_background(cell, HEADER_BLUE_HEX)
        
    sim_rows = [
        ("Adoção Clínico-Operacional (%)", "90%", "70%", "45%"),
        ("Retorno sobre Investimento (ROI)", "________ %", "________ %", "________ %"),
        ("Payback Período (Anos)", "________ anos", "________ anos", "________ anos"),
        ("Valor Presente Líquido (VPL)", "R$ ____________", "R$ ____________", "R$ ____________")
    ]
    
    for row_idx, data in enumerate(sim_rows):
        row = sim_table.rows[row_idx + 1]
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            cell.width = col_widths_sim[col_idx]
            p = cell.paragraphs[0]
            run = p.add_run(text)
            
            is_bold = col_idx > 0 and row_idx > 0
            color = BRAND_BLUE_RGB if is_bold else RGBColor(0,0,0)
            apply_text_styling(run, font_name="Arial", size_pt=9.5, bold=is_bold, color_rgb=color)
            
            if row_idx % 2 == 1:
                set_cell_background(cell, BG_LIGHT_HEX)
                
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    bp_p = doc.add_paragraph()
    bp_p.paragraph_format.line_spacing = 1.5
    bp_p.paragraph_format.first_line_indent = Cm(1.25)
    bp_run = bp_p.add_run("Ponto de Quebra Identificado: ")
    apply_text_styling(bp_run, font_name="Arial", size_pt=11, bold=True)
    bp_val = bp_p.add_run("[Cenário Pessimista / Cenário Realista / Viável em todos os cenários]\n")
    apply_text_styling(bp_val, font_name="Arial", size_pt=11)
    
    just_p = doc.add_paragraph()
    just_p.paragraph_format.line_spacing = 1.5
    just_p.paragraph_format.first_line_indent = Cm(1.25)
    just_run = just_p.add_run("Justificativa de Estresse:\n")
    apply_text_styling(just_run, font_name="Arial", size_pt=11, bold=True)
    just_val = just_p.add_run("[Insira aqui uma análise qualitativa sobre como a taxa de adoção e a ocorrência dos riscos ocultos afetam a viabilidade financeira e degradam as métricas de VPL e Payback...]")
    apply_text_styling(just_val, font_name="Arial", size_pt=11, italic=True, color_rgb=MUTED_SLATE_RGB)
    
    # --- 6. RECOMENDAÇÃO FINAL E GOVERNANÇA ---
    add_heading_abnt(doc, "6. Recomendação Final de Investimento (GO / NO-GO)", level=1)
    add_instruction_paragraph(doc, 
        "Instruções: Defina a recomendação formal da equipe para o C-level (GO para seguir com a contratação "
        "ou NO-GO para postergar/revisar o projeto). Fundamente sua recomendação cruzando o ROI ajustado ao estresse, "
        "a tolerância clínica aos riscos operacionais e o plano de governança traçado."
    )
    
    dec_p = doc.add_paragraph()
    dec_p.paragraph_format.line_spacing = 1.5
    dec_p.paragraph_format.first_line_indent = Cm(1.25)
    dec_run = dec_p.add_run("Recomendação Formal: ")
    apply_text_styling(dec_run, font_name="Arial", size_pt=11, bold=True)
    dec_val = dec_p.add_run("[   ] DECISÃO GO (APROVAÇÃO)    /    [   ] DECISÃO NO-GO (REJEIÇÃO)")
    apply_text_styling(dec_val, font_name="Arial", size_pt=11, bold=True, color_rgb=BRAND_BLUE_RGB)
    
    just_p2 = doc.add_paragraph()
    just_p2.paragraph_format.line_spacing = 1.5
    just_p2.paragraph_format.first_line_indent = Cm(1.25)
    just_lbl = just_p2.add_run("Justificativa de Governança:\n")
    apply_text_styling(just_lbl, font_name="Arial", size_pt=11, bold=True)
    just_val2 = just_p2.add_run(
        "[Fundamente a recomendação final de investimento. "
        "Explique por que os mitigantes reduzem os riscos ocultos a níveis aceitáveis "
        "ou por que a incerteza e a exposição ao passivo recomendam o NO-GO do projeto...]"
    )
    apply_text_styling(just_val2, font_name="Arial", size_pt=11, italic=True, color_rgb=MUTED_SLATE_RGB)

    # --- 7. REFERÊNCIAS ---
    add_heading_abnt(doc, "7. Referências Bibliográficas", level=1)
    add_instruction_paragraph(doc,
        "Instruções: Liste as referências utilizadas para fundamentar o relatório final em formato ABNT NBR 6023. "
        "Abaixo constam as referências básicas da disciplina para servir de guia para a equipe:"
    )
    
    # ABNT references do not have indentation on first line, but are spaced out
    ref_list = [
        "ARBACHE, Fernando. Inovação e Inteligência Artificial Aplicada à Gestão de Riscos. Porto Alegre: KnowRISK Press, 2024.",
        "CLEVERLEY, William O.; CLEVERLEY, James O.; SONG, Paula H. Essentials of Health Care Finance. 9. ed. Burlington: Jones & Bartlett Learning, 2023.",
        "DAMODARAN, Aswath. Avaliação de Investimentos: ferramentas e técnicas para a determinação do valor de qualquer ativo. 4. ed. Rio de Janeiro: GEN LTC, 2024.",
        "HUBBARD, Douglas W. How to Measure Anything: finding the value of intangibles in business. 4. ed. Hoboken: Wiley, 2024.",
        "TOPOL, Eric. Deep Medicine: how artificial intelligence can make healthcare human again. New York: Basic Books, 2019."
    ]
    
    for ref_str in ref_list:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # ABNT NBR 6023 style: single space, left aligned, but double spacing between references (accomplished via space_after)
        run = p.add_run(ref_str)
        apply_text_styling(run, font_name="Arial", size_pt=11)

    # Save the template to the public directory
    output_dir = "/Users/faila/Library/CloudStorage/GoogleDrive-failapereira@gmail.com/My Drive/MOINHOS/public"
    os.makedirs(output_dir, exist_ok=True)
    file_path = os.path.join(output_dir, "Trabalho_Final_Template.docx")
    doc.save(file_path)
    print(f"ABNT Word Template successfully generated at: {file_path}")

if __name__ == "__main__":
    main()
