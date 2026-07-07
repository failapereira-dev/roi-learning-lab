import os
import re
import json
import docx
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Colors
BRAND_BLUE_RGB = RGBColor(2, 132, 199)
SECONDARY_GRAY_RGB = RGBColor(59, 81, 102)
MUTED_SLATE_RGB = RGBColor(107, 124, 147)
BG_LIGHT_HEX = "F0F9FF"
HEADER_BLUE_HEX = "0284C7"
BORDER_GRAY_HEX = "CBD5E1"

def set_cell_background(cell, color_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def apply_text_styling(run, font_name="Arial", size_pt=12, bold=False, italic=False, color_rgb=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb
    else:
        run.font.color.rgb = RGBColor(0, 0, 0) # ABNT standard black

def add_heading_abnt(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.line_spacing = 1.5
    
    run = p.add_run(text)
    if level == 1:
        apply_text_styling(run, font_name="Arial", size_pt=14, bold=True)
    elif level == 2:
        apply_text_styling(run, font_name="Arial", size_pt=12, bold=True)
    elif level == 3:
        apply_text_styling(run, font_name="Arial", size_pt=12, bold=True, italic=True)
    return p

def add_body_paragraph(doc, text="", bold=False, italic=False, align_justify=True, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    if align_justify:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    
    if text:
        run = p.add_run(text)
        apply_text_styling(run, font_name="Arial", size_pt=12, bold=bold, italic=italic)
    return p

def add_body_paragraph_with_html(doc, html_text, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
        
    # Clean up some basic HTML tags
    clean_text = html_text.replace("<br>", "\n").replace("<br/>", "\n").replace("<br />", "\n")
    # Split by strong tags
    parts = re.split(r'(<strong>.*?</strong>)', clean_text)
    for part in parts:
        if part.startswith("<strong>") and part.endswith("</strong>"):
            clean_part = part[8:-9]
            run = p.add_run(clean_part)
            apply_text_styling(run, font_name="Arial", size_pt=12, bold=True)
        else:
            run = p.add_run(part)
            apply_text_styling(run, font_name="Arial", size_pt=12)
    return p

def add_list_item_abnt(doc, text, is_numbered=False):
    style = 'List Number' if is_numbered else 'List Bullet'
    p = doc.add_paragraph(style=style)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    
    run = p.add_run(text)
    apply_text_styling(run, font_name="Arial", size_pt=12)
    return p

def create_table_borders(table):
    tblPr = table._tbl.tblPr
    borders_xml = f'''
    <w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="{BORDER_GRAY_HEX}"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{BORDER_GRAY_HEX}"/>
        <w:left w:val="single" w:sz="4" w:space="0" w:color="{BORDER_GRAY_HEX}"/>
        <w:right w:val="single" w:sz="4" w:space="0" w:color="{BORDER_GRAY_HEX}"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{BORDER_GRAY_HEX}"/>
        <w:insideV w:val="none"/>
    </w:tblBorders>
    '''
    tblPr.append(parse_xml(borders_xml))

def clean_html_to_markdown(html_text):
    text = html_text
    text = text.replace("<strong>", "**").replace("</strong>", "**")
    text = text.replace("<br>", "\n").replace("<br/>", "\n").replace("<br />", "\n")
    return text

def main():
    # Load db_dump.json
    with open("db_dump.json", "r", encoding="utf-8") as f:
        db = json.load(f)
        
    scenarios = db["scenarios"]
    classes = db["classesContent"]
    
    # ------------------ GENERATE MARKDOWN GUIDE ------------------
    md_content = []
    
    md_content.append("# Guia Norteador do Professor: Avaliação de Riscos e Investimentos ROI")
    md_content.append("### MBA em Inteligência Artificial Aplicada à Saúde")
    md_content.append("#### Faculdade de Ciências da Saúde Moinhos de Vento\n")
    md_content.append("Este documento serve como roteiro pedagógico, gabarito de dinâmicas e barema oficial para a condução da disciplina. Ele foi estruturado para apoiar o professor na facilitação das três sessões interativas do portal de simulação, fornecendo o lastro conceitual e prático necessário.\n")
    md_content.append("---\n")
    
    # Section 1
    md_content.append("## 1. Roteiro e Cronograma Pedagógico Geral\n")
    md_content.append("A disciplina é composta por 3 aulas síncronas de 3 horas cada (totalizando 9 horas), com o seguinte cronograma sugerido de facilitação:\n")
    md_content.append("| Bloco de Tempo | Atividade | Papel do Facilitador |")
    md_content.append("| :--- | :--- | :--- |")
    md_content.append("| **00:00 - 00:20** (20 min) | Abertura & Icebreaker | Explicar dinâmica síncrona; comentar nuvem de palavras inicial. |")
    md_content.append("| **00:20 - 01:10** (50 min) | Bloco Teórico 1 & Enquetes | Explicar conceitos; disparar as perguntas e revelar as respostas corretas. |")
    md_content.append("| **01:10 - 01:40** (30 min) | Estudo de Caso & Leitura Rápida | Disparar a leitura síncrona (5 min) e debater as repercussões locais/globais. |")
    md_content.append("| **01:40 - 02:30** (50 min) | Dinâmica em Grupo Co-laborativa | Monitorar os workspaces das equipes e dar mentoria nos chats específicos. |")
    md_content.append("| **02:30 - 02:50** (20 min) | Debrief do Trabalho das Equipes | Projetar os resultados enviados e comentar as melhores análises. |")
    md_content.append("| **02:50 - 03:00** (10 min) | Reflexão Final & Encerramento | Consolidação do aprendizado e orientações para a próxima pré-aula.\n")
    md_content.append("---\n")
    
    # Section 2
    md_content.append("## 2. Guia de Respostas e Gabarito das Dinâmicas Síncronas\n")
    
    md_content.append("### AULA 1: Fundamentos de Risco em Inovação")
    md_content.append("* **Icebreaker (Tiro no Pé Tecnológico):**")
    md_content.append("  * *Pergunta:* \"Em uma palavra: qual foi a tecnologia que prometeu muito e entregou problema na saúde?\"")
    md_content.append("  * *Sugestões para a Nuvem de Palavras:* IBM Watson, Theranos, Chatbots, Prontuário sem triagem, IA de laudo genérica, Sensores sem fio, Reconhecimento de voz.")
    md_content.append("  * *Direcionamento:* O professor deve enfatizar que o fracasso raramente decorre de limitações da tecnologia isolada, mas sim da falta de integração de processos e falha em prever custos e riscos operacionais secundários.")
    md_content.append("* **Concept Check (Aula 1):**")
    md_content.append("  * *Pergunta 1 (Risco Oculto):* Qual das seguintes opções descreve melhor o conceito de \"Risco Oculto\" em projetos de IA na saúde?")
    md_content.append("    * **Resposta Correta:** A) Custos secundários de suporte técnico, retrabalho assistencial e processos por negligência decorrentes do uso da ferramenta.")
    md_content.append("  * *Pergunta 2 (Fadiga de Alarmes):* Qual o principal fator clínico que desloca um risco raro para o quadrante vermelho da matriz de riscos na saúde?")
    md_content.append("    * **Resposta Correta:** B) A gravidade do impacto clínico sobre a segurança do paciente, independente de sua probabilidade de ocorrência.")
    md_content.append("* **Classificação de Riscos ao Vivo (Slides 12a a 12d):**")
    md_content.append("  * *Risco Técnico:* Taxa de falsos negativos de 1% em pacientes com infarto agudo. (Gabarito: **Evitar / Mitigar** - Quadrante Vermelho devido à alta gravidade clínica).")
    md_content.append("  * *Risco Operacional:* Aceleração do tempo de ressonância gera superlotação de recepção e estacionamento. (Gabarito: **Mitigar / Monitorar** - Quadrante Laranja/Amarelo).")
    md_content.append("  * *Risco Clínico-Cultural:* Equipe de enfermagem ignora alertas de arritmia por fadiga de alarmes. (Gabarito: **Evitar / Mitigar** - Quadrante Vermelho pelo potencial catastrófico).")
    md_content.append("  * *Risco Financeiro:* Necessidade de upgrade milionário de Wi-Fi para telemetria cardíaca. (Gabarito: **Mitigar / Monitorar** - Quadrante Laranja/Amarelo).\n")
    
    md_content.append("### AULA 2: Mitigação de Risco e a Anatomia do ROI Inteligente")
    md_content.append("* **Check-in (O Custo Oculto que Ninguém Somou):**")
    md_content.append("  * *Pergunta:* \"Qual custo mais 'some' das planilhas de projetos na saúde?\"")
    md_content.append("  * *Opções / Resposta Sugerida:* **Treinamento e Adoção** (ou **Retrabalho/Ineficiência**). A adoção inadequada é o principal fator de depreciação do ROI na saúde.")
    md_content.append("* **Concept Check (Aula 2):**")
    md_content.append("  * *Pergunta 1 (ROI sob Estresse):* Por que realizamos o estresse de adoção operacional no cálculo do ROI?")
    md_content.append("    * **Resposta Correta:** B) Porque o retorno financeiro teórico depende diretamente de o corpo clínico efetivamente utilizar a ferramenta de IA no fluxo de trabalho real.")
    md_content.append("  * *Pergunta 2 (Custo de Riscos Ocultos):* Como o \"Custo de Riscos Mapeados\" afeta a viabilidade financeira de um projeto de IA?")
    md_content.append("    * **Resposta Correta:** A) Ele atua como um custo operacional anualizado que reduz diretamente o fluxo de caixa líquido e afeta o VPL.\n")
    
    md_content.append("### AULA 3: Governança, Tomada de Decisão e Business Case")
    md_content.append("* **Check-in (O que trava a Diretoria):**")
    md_content.append("  * *Pergunta:* \"Em uma palavra: o que mais trava a aprovação de um bom projeto na diretoria?\"")
    md_content.append("  * *Sugestões para a Nuvem de Palavras:* CAPEX, Payback longo, Falta de Dados, Insegurança, Desalinhamento, C-Level, ROI abstrato.")
    md_content.append("* **Concept Check (Aula 3):**")
    md_content.append("  * *Pergunta 1 (Business Case Executivo):* O que a diretoria executiva (C-level) mais busca em um Business Case de IA na saúde?")
    md_content.append("    * **Resposta Correta:** C) A demonstração da viabilidade financeira realista sob condições de estresse combinada com garantias de segurança ao paciente.")
    md_content.append("  * *Pergunta 2 (Matriz Ponderada):* Como a Matriz de Decisão Ponderada auxilia na governança corporativa?")
    md_content.append("    * **Resposta Correta:** B) Permitindo avaliar o projeto sob critérios múltiplos (financeiro, segurança, facilidade técnica), evitando a dependência de uma única métrica (ROI simples).\n")
    md_content.append("---\n")
    
    # Section 3
    md_content.append("## 3. Detalhamento dos Cenários Narrativos das Equipes\n")
    md_content.append("Os cenários foram reescritos em formato narrativo para ocultar a estrutura direta de respostas. A tabela abaixo fornece a transcrição de cada caso para suporte do professor, mapeando os parâmetros numéricos e de riscos subjacentes que as equipes devem extrair:\n")
    
    # Group 1 details
    md_content.append("### Grupo 1: IA de Triagem de Urgência")
    md_content.append(f"* **Transcrição do Caso:** {clean_html_to_markdown(scenarios['1']['caseStudy'])}")
    md_content.append("""* **Parâmetros Financeiros Reais:**
  * Investimento Inicial (CAPEX): R$ 1.500.000
  * Benefício Anual (OPEX Retorno): R$ 300.000
  * Custo Anual de Manutenção: R$ 50.000
  * Custo Anual de Riscos Mapeados (Matriz): R$ 350.000
  * Taxa de Adoção: Esperada (90%), Realista (70%), Pessimista (45%)
* **Diretrizes de Mapeamento de Riscos (Aula 1):**
  * *Risco Técnico:* Taxa de falsos negativos que trie infarto como verde. Mitigação: Dupla checagem obrigatória pelo enfermeiro de triagem e travas eletrônicas para queixas cardíacas.
  * *Risco Operacional:* Fadiga de alarmes gerando cliques extras e rejeição dos triadores. Mitigação: Simplificação do design da interface (UX) e parametrização dos alarmes de acordo com o perfil do hospital.
  * *Risco Clínico-Cultural:* Resistência dos pacientes em serem triados por algoritmos. Mitigação: Campanhas de comunicação interna e acolhimento humano na entrega das pulseiras.
  * *Risco Financeiro:* Custos de processos judiciais de R$ 350.000. Mitigação: Seguro de responsabilidade civil médica e auditoria constante de acurácia clínica da IA.
  * *Fator Showstopper:* A acurácia do software na detecção de queixas de risco de morte imediata (como infarto ou sepse).
* **Análise da Simulação de ROI (Aula 2):**
  * O projeto enfrenta degradação imediata de ROI quando os custos de riscos ocultos (R$ 350.000) superam o retorno operacional líquido. O ponto de quebra financeira é atingido no Cenário Pessimista (onde a adoção de 45% anula os benefícios de vazão).
* **Diretrizes para o Relatório Final (Aula 3):**
  * A equipe deve decidir por **GO com mitigadores estritos** ou **NO-GO provisório** até que a taxa de falsos negativos caia. A fundamentação deve cruzar o custo do processo judicial (R$ 350K) com o benefício total.
""")
    md_content.append("---\n")
    
    # Group 2 details
    md_content.append("### Grupo 2: Automação de Farmácia Robótica")
    md_content.append(f"* **Transcrição do Caso:** {clean_html_to_markdown(scenarios['2']['caseStudy'])}")
    md_content.append("""* **Parâmetros Financeiros Reais:**
  * Investimento Inicial (CAPEX): R$ 1.800.000
  * Benefício Anual (OPEX Retorno): R$ 320.000
  * Custo Anual de Manutenção: R$ 60.000
  * Custo Anual de Riscos Mapeados (Matriz): R$ 400.000
  * Taxa de Adoção: Esperada (90%), Realista (70%), Pessimista (45%)
* **Diretrizes de Mapeamento de Riscos (Aula 1):**
  * *Risco Técnico:* Erro de leitura de prescrição legada propagando erros do robô. Mitigação: Camada de integração (middleware) robusta com validação em tempo real de consistência de dosagem.
  * *Risco Operacional:* Pane mecânica total paralisando farmácia. Mitigação: Plano de redundância com dispensador secundário mecânico manual de backup e contrato de SLA de 2 horas.
  * *Risco Clínico-Cultural:* Resistência dos farmacêuticos. Mitigação: Treinamentos práticos, foco em ergonomia e posicionamento do farmacêutico em farmácia clínica de leito.
  * *Risco Financeiro:* Custo de contingência emergencial de R$ 400.000. Mitigação: Seguro patrimonial de operação contínua e provisionamento financeiro prévio.
  * *Fator Showstopper:* A integridade mecânica e conectividade de rede do robô.
* **Análise da Simulação de ROI (Aula 2):**
  * Ponto de Quebra: O projeto se torna altamente inviável sob estresse (Cenário Pessimista e Realista), pois a economia de R$ 320.000 é anulada pelo custo recorrente de manutenção (R$ 60.000) somado ao risco potencializado da contingência emergencial.
* **Diretrizes para o Relatório Final (Aula 3):**
  * A recomendação natural é de **NO-GO** ou **Aprovação Condicional** com renegociação do CAPEX. A justificativa deve focar no alto payback inicial decorrente da baixa margem de benefício operacional simples.
""")
    md_content.append("---\n")

    # Group 3 details
    md_content.append("### Grupo 3: Telemetria Cardíaca Wireless")
    md_content.append(f"* **Transcrição do Caso:** {clean_html_to_markdown(scenarios['3']['caseStudy'])}")
    md_content.append("""* **Parâmetros Financeiros Reais:**
  * Investimento Inicial (CAPEX): R$ 2.200.000
  * Benefício Anual (OPEX Retorno): R$ 1.500.000
  * Custo Anual de Manutenção: R$ 180.000
  * Custo Anual de Riscos Mapeados (Matriz): R$ 800.000
  * Taxa de Adoção: Esperada (90%), Realista (70%), Pessimista (45%)
* **Diretrizes de Mapeamento de Riscos (Aula 1):**
  * *Risco Técnico:* Oscilação de rede Wi-Fi hospitalar gerando perda de monitoramento contínuo. Mitigação: Mapeamento de calor de sinal e instalação de access points redundantes na enfermaria.
  * *Risco Operacional:* Fadiga de alarmes por ruído físico. Mitigação: Parametrizar a IA para ignorar artefatos de movimento físico do paciente e escalonamento inteligente.
  * *Risco Clínico-Cultural:* Enfermagem desligar os smartphones de alerta. Mitigação: Dashboard central visível no posto e regras rígidas de conformidade assistencial.
  * *Risco Financeiro:* Processos judiciais por atraso assistencial estimados em R$ 800.000. Mitigação: Seguro de responsabilidade civil médico e auditoria constante de acurácia de alertas.
  * *Fator Showstopper:* Conectividade de rede Wi-Fi e estabilidade da telemetria móvel.
* **Análise da Simulação de ROI (Aula 2):**
  * O projeto possui excelente viabilidade financeira devido ao alto retorno projetado de R$ 1,5 milhão (pela liberação acelerada de leitos de enfermaria). Ele suporta o teste de estresse realista e pessimista, mas exige mitigação rígida do Wi-Fi para evitar o passivo de R$ 800.000.
* **Diretrizes para o Relatório Final (Aula 3):**
  * Recomendação de **GO (Aprovado)**. A justificativa deve estruturar o business case demonstrando que, mesmo sob cenário de estresse pessimista, a telemetria se paga em poucos anos, com foco estratégico em segurança do paciente pós-infarto.
""")
    md_content.append("---\n")

    # Group 4 details
    md_content.append("### Grupo 4: IA em Radiologia")
    md_content.append(f"* **Transcrição do Caso:** {clean_html_to_markdown(scenarios['4']['caseStudy'])}")
    md_content.append("""* **Parâmetros Financeiros Reais:**
  * Investimento Inicial (CAPEX): R$ 800.000
  * Benefício Anual (OPEX Retorno): R$ 80.000.000
  * Custo Anual de Manutenção: R$ 150.000
  * Custo Anual de Riscos Mapeados (Matriz): R$ 50.000.000
  * Taxa de Adoção: Esperada (90%), Realista (70%), Pessimista (45%)
* **Diretrizes de Mapeamento de Riscos (Aula 1):**
  * *Risco Técnico:* Perda de acurácia em pediatria ou equipamentos legados. Mitigação: Treinamento específico da IA com dados locais e exclusão temporária de laudos automáticos para menores de 12 anos.
  * *Risco Operacional:* Sobrecarga por revisão de falsos positivos. Mitigação: Calibragem dinâmica da sensibilidade e especificidade da IA junto ao fornecedor.
  * *Risco Clínico-Cultural:* Confiança cega do radiologista (Over-reliance). Mitigação: Proibição de laudar sem a visualização prévia e completa da imagem original (IA cega na primeira leitura).
  * *Risco Financeiro:* Custo de risco judicial extraordinário de R$ 50.000.000. Mitigação: Seguro de responsabilidade corporativa médico de alta cobertura e auditoria semanal amostral de laudos.
  * *Fator Showstopper:* A dependência excessiva dos radiologistas seniores nas sugestões da IA.
* **Análise da Simulação de ROI (Aula 2):**
  * Embora o benefício econômico de escala seja massivo (R$ 80M), o passivo judicial de R$ 50M atua como um enorme detrator de ROI se a mitigação de confiança falhar. O projeto tem viabilidade extrema no cenário esperado, mas no cenário realista com alta sinistralidade o VPL é significativamente corroído.
* **Diretrizes para o Relatório Final (Aula 3):**
  * Recomendação de **GO condicionado à implantação de fluxo 'Double-Blind'**. O business case deve detalhar as travas assistenciais para impedir a dependência cega antes de liberar o CAPEX.
""")
    md_content.append("---\n")

    # Group 5 details
    md_content.append("### Grupo 5: Prontuário Eletrônico Integrado")
    md_content.append(f"* **Transcrição do Caso:** {clean_html_to_markdown(scenarios['5']['caseStudy'])}")
    md_content.append("""* **Parâmetros Financeiros Reais:**
  * Investimento Inicial (CAPEX): R$ 5.000.000
  * Benefício Anual (OPEX Retorno): R$ 1.050.000
  * Custo Anual de Manutenção: R$ 450.000
  * Custo Anual de Riscos Mapeados (Matriz): R$ 2.000.000
  * Taxa de Adoção: Esperada (90%), Realista (70%), Pessimista (45%)
* **Diretrizes de Mapeamento de Riscos (Aula 1):**
  * *Risco Técnico:* Perda de dados históricos importantes de migração. Mitigação: ETL pré-validado em ambiente de staging e cópia física espelhada de histórico médico em banco de leitura.
  * *Risco Operacional:* Pane ou paralisação na virada de chave (Go-Live). Mitigação: Realizar a transição em feriado/fim de semana e manter equipes médicas em escala de plantão ampliada por 72 horas.
  * *Risco Clínico-Cultural:* Rejeição de médicos por lentidão de digitação. Mitigação: Copiloto de IA de voz ativa e simplificação de templates de prescrição padrão.
  * *Risco Financeiro:* Custos de contingência de go-live de R$ 2.000.000. Mitigação: Escalonamento financeiro em contrato de transição com o fornecedor do prontuário.
  * *Fator Showstopper:* A integridade histórica dos registros médicos dos pacientes.
* **Análise da Simulação de ROI (Aula 2):**
  * Com o alto CAPEX (R$ 5,0M) e a manutenção anual pesada (R$ 450K), o retorno de R$ 1,05M de redução de glosas tem payback lento. A adição do risco de pane no go-live (R$ 2,0M) empurra o projeto para o prejuízo (VPL negativo) em cenários de estresse de adoção realista ou pessimista.
* **Diretrizes para o Relatório Final (Aula 3):**
  * A recomendação pode ser **NO-GO** devido ao alto risco operacional ou **GO Faseado** (por exemplo, implantar primeiro apenas na UTI e faturamento). A discussão de governança deve pesar o risco de manter sistemas desintegrados vs. os custos da transição.
""")
    md_content.append("---\n")
    
    # Section 4
    md_content.append("## 4. Barema para Correção do Trabalho Final (10 Pontos)\n")
    md_content.append("O trabalho final consiste no envio por e-mail do Relatório Executivo (.docx) gerado e completado por cada equipe. O professor utilizará as diretrizes abaixo para atribuir notas de 0 a 10 pontos:\n")
    md_content.append("| Critério | Peso | Descrição do Nível Excelente (2,0 Pontos) | Descrição do Nível Regular (1,0 Ponto) | Descrição do Nível Insuficiente (0 Pontos) |")
    md_content.append("| :--- | :--- | :--- | :--- | :--- |")
    md_content.append("| **Conexão Contexto-Problema** | 2,0 pt | Extrai corretamente todos os dados de baseline do caso (volumetria de exames/dispensações, erros atuais de triagem/glosas) e descreve o gargalo clínico de forma precisa e integrada. | Descreve o problema em linhas gerais, mas falha em quantificar de forma exata a volumetria e os prejuízos iniciais descritos no cenário. | Não contextualiza o cenário da instituição ou apresenta dados inventados não relacionados ao grupo de trabalho correspondente. |")
    md_content.append("| **Modelagem e Mitigação de Riscos** | 2,0 pt | Identifica e classifica com rigor os 4 tipos de riscos (Técnico, Operacional, Clínico-Cultural e Financeiro) e propõe planos de mitigação práticos, realistas e específicos para a saúde. | Lista os riscos de forma superficial ou propõe mitigantes genéricos (ex: \"comprar melhor antivírus\" ou \"dar mais treinamento\") sem detalhar o fluxo. | Deixa de analisar uma ou mais categorias de riscos ou não propõe estratégias de contingência assistencial/tecnológica relevantes. |")
    md_content.append("| **Simulação Financeira sob Estresse** | 2,0 pt | Transcreve e analisa matematicamente de forma exata os valores de ROI, Payback e VPL para os 3 cenários (esperado, realista e pessimista), apontando corretamente o ponto de quebra. | Preenche a tabela financeira, mas comete erros leves de transcrição ou falha em justificar a degradação das métricas sob o teste de estresse de adoção. | A tabela financeira apresenta dados em branco, erros grosseiros de cálculo ou ausência completa de justificativa analítica do estresse. |")
    md_content.append("| **Recomendação e Governança** | 2,0 pt | Apresenta recomendação clara de GO/NO-GO fundamentada estrategicamente. Cruza a tolerância assistencial ao risco clínico com a viabilidade econômica de longo prazo da instituição. | Define a decisão de GO ou NO-GO, mas a justificativa é rasa, puramente baseada em intuição subjetiva ou com fraco alinhamento aos dados financeiros sob estresse. | Ausência de decisão de recomendação de investimento formalizada ou falta absoluta de justificativa baseada em governança clínica e corporativa. |")
    md_content.append("| **Qualidade do Relatório Executivo** | 2,0 pt | O relatório final está integralmente preenchido, formatado sem os textos de instruções do template, apresentando clareza de linguagem executiva para C-Level e rigor técnico. | Relatório preenchido, mas mantém trechos das instruções de preenchimento, apresenta erros de digitação ou linguagem pouco executiva/inadequada para diretoria. | Relatório incompleto, desorganizado, sem o preenchimento de seções obrigatórias ou enviado de forma desestruturada.\n")
    md_content.append("---\n")

    # Section 5 (Dynamic slides and notes)
    md_content.append("## 5. Roteiro de Condução e Notas de Facilitador Slide a Slide\n")
    
    for c in classes:
        class_id = c["id"]
        class_title = c["title"]
        md_content.append(f"### {class_title}")
        for s in c["slides"]:
            slide_id = s["id"]
            slide_title = s["title"]
            raw_notes = s.get("notes", "Nenhuma nota de condução definida.")
            clean_notes = clean_html_to_markdown(raw_notes)
            md_content.append(f"* **{slide_id} ({slide_title}):** {clean_notes}")
        md_content.append("") # newline between classes
        
    # Save Guia_Norteador_Professor.md
    with open("Guia_Norteador_Professor.md", "w", encoding="utf-8") as f:
        f.write("\n".join(md_content))
        
    print("Guia_Norteador_Professor.md successfully generated!")

    # ------------------ GENERATE WORD GUIDE ------------------
    doc = docx.Document()
    
    # Margins ABNT
    for section in doc.sections:
        section.top_margin = Cm(3.0)
        section.left_margin = Cm(3.0)
        section.bottom_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        
    # --- CAPA ---
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_p.paragraph_format.space_after = Pt(12)
    if os.path.exists("Logo Moinhos.png"):
        logo_p.add_run().add_picture("Logo Moinhos.png", width=Inches(1.8))
        
    inst_p = doc.add_paragraph()
    inst_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inst_p.paragraph_format.space_after = Pt(72)
    inst_p.paragraph_format.line_spacing = 1.15
    run_inst1 = inst_p.add_run("FACULDADE DE CIÊNCIAS DA SAÚDE MOINHOS DE VENTO\n")
    apply_text_styling(run_inst1, font_name="Arial", size_pt=12, bold=True)
    run_inst2 = inst_p.add_run("MBA EM INTELIGÊNCIA ARTIFICIAL APLICADA À SAÚDE\nDISCIPLINA DE AVALIAÇÃO DE RISCOS E INVESTIMENTOS ROI")
    apply_text_styling(run_inst2, font_name="Arial", size_pt=11)
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(12)
    run_title = title_p.add_run("DOCUMENTO NORTEADOR DO PROFESSOR")
    apply_text_styling(run_title, font_name="Arial", size_pt=16, bold=True)
    
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(84)
    run_sub = sub_p.add_run("Guia Pedagógico Completo, Speaker Notes, Gabarito das Dinâmicas e Barema de Correção")
    apply_text_styling(run_sub, font_name="Arial", size_pt=12, italic=True, color_rgb=SECONDARY_GRAY_RGB)
    
    city_p = doc.add_paragraph()
    city_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    city_p.paragraph_format.space_before = Pt(120)
    city_p.paragraph_format.space_after = Pt(0)
    run_city = city_p.add_run("PORTO ALEGRE\n2026")
    apply_text_styling(run_city, font_name="Arial", size_pt=12, bold=True)
    
    doc.add_page_break()
    
    # --- 1. INTRODUÇÃO E CRONOGRAMA ---
    add_heading_abnt(doc, "1. Introdução e Cronograma Pedagógico Geral", level=1)
    add_body_paragraph(doc, 
        "Este documento serve como roteiro pedagógico, gabarito de dinâmicas e barema oficial para a condução da disciplina. "
        "Ele foi estruturado para apoiar o professor na facilitação das três sessões interativas do portal de simulação, "
        "fornecendo o lastro conceitual e prático necessário para avaliar a maturidade das equipes nos testes de estresse operacional e análise de ROI."
    )
    
    add_body_paragraph(doc, "Abaixo, apresenta-se o cronograma sugerido para cada aula síncrona de 3 horas:", bold=True, indent=False)
    
    # Agenda Table
    agenda_table = doc.add_table(rows=7, cols=3)
    agenda_table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    create_table_borders(agenda_table)
    
    headers_agenda = ["Bloco de Tempo", "Atividade", "Papel do Facilitador"]
    col_widths_agenda = [Cm(3.5), Cm(5.0), Cm(7.5)]
    
    for idx, text in enumerate(headers_agenda):
        cell = agenda_table.rows[0].cells[idx]
        cell.width = col_widths_agenda[idx]
        p = cell.paragraphs[0]
        run = p.add_run(text)
        apply_text_styling(run, font_name="Arial", size_pt=10, bold=True, color_rgb=RGBColor(255, 255, 255))
        set_cell_background(cell, HEADER_BLUE_HEX)
        
    agenda_rows = [
        ("00:00 - 00:20 (20 min)", "Abertura & Icebreaker", "Explicar dinâmica síncrona; comentar nuvem de palavras inicial."),
        ("00:20 - 01:10 (50 min)", "Bloco Teórico 1 & Enquetes", "Explicar conceitos; disparar as perguntas e revelar as respostas corretas."),
        ("01:10 - 01:40 (30 min)", "Estudo de Caso & Leitura Rápida", "Disparar a leitura síncrona (5 min) e debater as repercussões locais/globais."),
        ("01:40 - 02:30 (50 min)", "Dinâmica em Grupo Co-laborativa", "Monitorar os workspaces das equipes e dar mentoria nos chats específicos."),
        ("02:30 - 02:50 (20 min)", "Debrief do Trabalho das Equipes", "Projetar os resultados enviados e comentar as melhores análises."),
        ("02:50 - 03:00 (10 min)", "Reflexão Final & Encerramento", "Consolidação do aprendizado e orientações para a próxima pré-aula.")
    ]
    
    for row_idx, data in enumerate(agenda_rows):
        row = agenda_table.rows[row_idx + 1]
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            cell.width = col_widths_agenda[col_idx]
            p = cell.paragraphs[0]
            run = p.add_run(text)
            apply_text_styling(run, font_name="Arial", size_pt=9.5, color_rgb=SECONDARY_GRAY_RGB)
            if row_idx % 2 == 1:
                set_cell_background(cell, BG_LIGHT_HEX)
                
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # --- 2. GABARITO DAS DINÂMICAS ---
    add_heading_abnt(doc, "2. Guia de Respostas e Gabarito das Dinâmicas Síncronas", level=1)
    
    add_heading_abnt(doc, "Aula 1: Fundamentos de Risco em Inovação", level=2)
    add_body_paragraph(doc, "• Icebreaker (Tiro no Pé): Pense em uma tecnologia que prometeu muito e entregou problema. Sugestões de nuvem: IBM Watson, Theranos, Chatbots, Prontuário sem triagem, IA de laudo genérica. Direcionamento: O risco quase nunca está na tecnologia, mas sim em como ela se integra aos fluxos e pessoas.")
    add_body_paragraph(doc, "• Concept Check 1: Risco Oculto (Resp: Custos secundários de suporte, retrabalho e passivos judiciais) e Matriz de Riscos na Saúde (Resp: A gravidade do impacto clínico sobre o paciente dita a gravidade, mesmo sendo um evento raro).")
    add_body_paragraph(doc, "• Enquetes de Riscos ao Vivo: Risco Clínico (Taxa de 1% falsos negativos infarto -> Evitar/Mitigar, Quadrante Vermelho); Risco Operacional (Saturação de recepção -> Mitigar/Monitorar, Quadrante Laranja/Amarelo).")
    
    add_heading_abnt(doc, "Aula 2: Mitigação de Risco e a Anatomia do ROI Inteligente", level=2)
    add_body_paragraph(doc, "• Check-in: Custo invisível mais negligenciado em propostas. Sugestão: Treinamento e Adoção. O professor deve pontuar que a falta de engajamento do time assistencial anula os ganhos projetados na planilha teórica.")
    add_body_paragraph(doc, "• Concept Check 2: Estresse de Adoção (Resp: O retorno financeiro real depende diretamente de o corpo clínico efetivamente utilizar a ferramenta de IA no fluxo de trabalho real) e Custo de Riscos Ocultos (Resp: Atua como um detrator de fluxo de caixa líquido anualizado, degradando o VPL).")
    
    add_heading_abnt(doc, "Aula 3: Governança, Tomada de Decisão e Business Case", level=2)
    add_body_paragraph(doc, "• Check-in: O que mais trava a aprovação na diretoria. Sugestões de nuvem: CAPEX alto, Payback longo, Falta de dados, Insegurança de risco regulatório.")
    add_body_paragraph(doc, "• Concept Check 3: Foco da diretoria (Resp: Viabilidade financeira em cenários realistas de estresse sob segurança assistencial) e Matriz Ponderada (Resp: Permite avaliar projetos sob critérios múltiplos complexos além do ROI financeiro simples).")
    
    doc.add_page_break()

    # --- 3. DETALHAMENTO DOS CASOS E PARÂMETROS ---
    add_heading_abnt(doc, "3. Detalhamento dos Cenários Narrativos das Equipes", level=1)
    add_body_paragraph(doc, "Os cenários foram reescritos em formato narrativo para ocultar a estrutura direta de respostas. Abaixo apresenta-se a transcrição integral e exata de cada caso para suporte do professor, mapeando os parâmetros numéricos e de riscos subjacentes que as equipes devem extrair:")
    
    # Replicate scenarios
    group_cases_details = [
        ("1", "Grupo 1: IA de Triagem de Urgência", 
         """* **Parâmetros Financeiros Reais:**
  • Investimento Inicial (CAPEX): R$ 1.500.000
  • Benefício Anual (OPEX Retorno): R$ 300.000
  • Custo Anual de Manutenção: R$ 50.000
  • Custo Anual de Riscos Mapeados (Matriz): R$ 350.000
  • Taxa de Adoção: Esperada (90%), Realista (70%), Pessimista (45%)
* **Diretrizes de Mapeamento de Riscos (Aula 1):**
  • Risco Técnico: Taxa de falsos negativos que trie infarto como verde. Mitigação: Dupla checagem obrigatória pelo enfermeiro de triagem e travas eletrônicas para queixas cardíacas.
  • Risco Operacional: Fadiga de alarmes gerando cliques extras e rejeição dos triadores. Mitigação: Simplificação do design da interface (UX) e parametrização dos alarmes de acordo com o perfil do hospital.
  • Risco Clínico-Cultural: Resistência dos pacientes em serem triados por algoritmos. Mitigação: Campanhas de comunicação interna e acolhimento humano na entrega das pulseiras.
  • Risco Financeiro: Custos de processos judiciais de R$ 350.000. Mitigação: Seguro de responsabilidade civil médica e auditoria constante de acurácia clínica da IA.
  • Fator Showstopper: A acurácia do software na detecção de queixas de risco de morte imediata (como infarto ou sepse).
* **Análise da Simulação de ROI (Aula 2):**
  • O projeto enfrenta degradação imediata de ROI quando os custos de riscos ocultos (R$ 350.000) superam o retorno operacional líquido. O ponto de quebra financeira é atingido no Cenário Pessimista (onde a adoção de 45% anula os benefícios de vazão).
* **Diretrizes para o Relatório Final (Aula 3):**
  • A equipe deve decidir por GO com mitigadores estritos ou NO-GO provisório até que a taxa de falsos negativos caia. A fundamentação deve cruzar o custo do processo judicial (R$ 350K) com o benefício total."""),
         
        ("2", "Grupo 2: Automação de Farmácia Robótica",
         """* **Parâmetros Financeiros Reais:**
  • Investimento Inicial (CAPEX): R$ 1.800.000
  • Benefício Anual (OPEX Retorno): R$ 320.000
  • Custo Anual de Manutenção: R$ 60.000
  • Custo Anual de Riscos Mapeados (Matriz): R$ 400.000
  • Taxa de Adoção: Esperada (90%), Realista (70%), Pessimista (45%)
* **Diretrizes de Mapeamento de Riscos (Aula 1):**
  • Risco Técnico: Erro de leitura de prescrição legada propagando erros do robô. Mitigação: Camada de integração (middleware) robusta com validação em tempo real de consistência de dosagem.
  • Risco Operacional: Pane mecânica total paralisando farmácia. Mitigação: Plano de redundância com dispensador secundário mecânico manual de backup e contrato de SLA de 2 horas.
  • Risco Clínico-Cultural: Resistência dos farmacêuticos. Mitigação: Treinamentos práticos, foco em ergonomia e posicionamento do farmacêutico em farmácia clínica de leito.
  • Risco Financeiro: Custo de contingência emergencial de R$ 400.000. Mitigação: Seguro patrimonial de operação contínua e provisionamento financeiro prévio.
  • Fator Showstopper: A integridade mecânica e conectividade de rede do robô.
* **Análise da Simulação de ROI (Aula 2):**
  • Ponto de Quebra: O projeto se torna altamente inviável sob estresse (Cenário Pessimista e Realista), pois a economia de R$ 320.000 é anulada pelo custo recorrente de manutenção (R$ 60.000) somado ao risco potencializado da contingência emergencial.
* **Diretrizes para o Relatório Final (Aula 3):**
  • A recomendação natural é de NO-GO ou Aprovação Condicional com renegociação do CAPEX. A justificativa deve focar no alto payback inicial decorrente da baixa margem de benefício operacional simples."""),
         
        ("3", "Grupo 3: Telemetria Cardíaca Wireless",
         """* **Parâmetros Financeiros Reais:**
  • Investimento Inicial (CAPEX): R$ 2.200.000
  • Benefício Anual (OPEX Retorno): R$ 1.500.000
  • Custo Anual de Manutenção: R$ 180.000
  • Custo Anual de Riscos Mapeados (Matriz): R$ 800.000
  • Taxa de Adoção: Esperada (90%), Realista (70%), Pessimista (45%)
* **Diretrizes de Mapeamento de Riscos (Aula 1):**
  • Risco Técnico: Oscilação de rede Wi-Fi hospitalar gerando perda de monitoramento contínuo. Mitigação: Mapeamento de calor de sinal e instalação de access points redundantes na enfermaria.
  • Risco Operacional: Fadiga de alarmes por ruído físico. Mitigação: Parametrizar a IA para ignorar artefatos de movimento físico do paciente e escalonamento inteligente.
  • Risco Clínico-Cultural: Enfermagem desligar os smartphones de alerta. Mitigação: Dashboard central visível no posto e regras rígidas de conformidade assistencial.
  • Risco Financeiro: Processos judiciais por atraso assistencial estimados em R$ 800.000. Mitigação: Seguro de responsabilidade civil médico e auditoria constante de acurácia de alertas.
  • Fator Showstopper: Conectividade de rede Wi-Fi e estabilidade da telemetria móvel.
* **Análise da Simulação de ROI (Aula 2):**
  • O projeto possui excelente viabilidade financeira devido ao alto retorno projetado de R$ 1,5 milhão (pela liberação acelerada de leitos de enfermaria). Ele suporta o teste de estresse realista e pessimista, mas exige mitigação rígida do Wi-Fi para evitar o passivo de R$ 800.000.
* **Diretrizes para o Relatório Final (Aula 3):**
  • Recomendação de GO (Aprovado). A justificativa deve estruturar o business case demonstrando que, mesmo sob cenário de estresse pessimista, a telemetria se paga em poucos anos, com foco estratégico em segurança do paciente pós-infarto."""),
         
        ("4", "Grupo 4: IA em Radiologia",
         """* **Parâmetros Financeiros Reais:**
  • Investimento Inicial (CAPEX): R$ 800.000
  • Benefício Anual (OPEX Retorno): R$ 80.000.000
  • Custo Anual de Manutenção: R$ 150.000
  • Custo Anual de Riscos Mapeados (Matriz): R$ 50.000.000
  • Taxa de Adoção: Esperada (90%), Realista (70%), Pessimista (45%)
* **Diretrizes de Mapeamento de Riscos (Aula 1):**
  • Risco Técnico: Perda de acurácia em pediatria ou equipamentos legados. Mitigação: Treinamento específico da IA com dados locais e exclusão temporária de laudos automáticos para menores de 12 anos.
  • Risco Operacional: Sobrecarga por revisão de falsos positivos. Mitigação: Calibragem dinâmica da sensibilidade e especificidade da IA junto ao fornecedor.
  • Risco Clínico-Cultural: Confiança cega do radiologista (Over-reliance). Mitigação: Proibição de laudar sem a visualização prévia e completa da imagem original (IA cega na primeira leitura).
  • Risco Financeiro: Custo de risco judicial extraordinário de R$ 50.000.000. Mitigação: Seguro de responsabilidade corporativa médico de alta cobertura e auditoria semanal amostral de laudos.
  • Fator Showstopper: A dependência excessiva dos radiologistas seniores nas sugestões da IA.
* **Análise da Simulação de ROI (Aula 2):**
  • Embora o benefício econômico de escala seja massivo (R$ 80M), o passivo judicial de R$ 50M atua como um enorme detrator de ROI se a mitigação de confiança falhar. O projeto tem viabilidade extrema no cenário esperado, mas no cenário realista com alta sinistralidade o VPL é significativamente corroído.
* **Diretrizes para o Relatório Final (Aula 3):**
  • Recomendação de GO condicionado à implantação de fluxo 'Double-Blind'. O business case deve detalhar as travas assistenciais para impedir a dependência cega antes de liberar o CAPEX."""),
         
        ("5", "Grupo 5: Prontuário Eletrônico Integrado",
         """* **Parâmetros Financeiros Reais:**
  • Investimento Inicial (CAPEX): R$ 5.000.000
  • Benefício Anual (OPEX Retorno): R$ 1.050.000
  • Custo Anual de Manutenção: R$ 450.000
  • Custo Anual de Riscos Mapeados (Matriz): R$ 2.000.000
  • Taxa de Adoção: Esperada (90%), Realista (70%), Pessimista (45%)
* **Diretrizes de Mapeamento de Riscos (Aula 1):**
  • Risco Técnico: Perda de dados históricos importantes de migração. Mitigação: ETL pré-validado em ambiente de staging e cópia física espelhada de histórico médico em banco de leitura.
  • Risco Operacional: Pane ou paralisação na virada de chave (Go-Live). Mitigação: Realizar a transição em feriado/fim de semana e manter equipes médicas em escala de plantão ampliada por 72 horas.
  • Risco Clínico-Cultural: Rejeição de médicos por lentidão de digitação. Mitigação: Copiloto de IA de voz ativa e simplificação de templates de prescrição padrão.
  • Risco Financeiro: Custos de contingência de go-live de R$ 2.000.000. Mitigação: Escalonamento financeiro em contrato de transição com o fornecedor do prontuário.
  • Fator Showstopper: A integridade histórica dos registros médicos dos pacientes.
* **Análise da Simulação de ROI (Aula 2):**
  • Com o alto CAPEX (R$ 5,0M) e a manutenção anual pesada (R$ 450K), o retorno de R$ 1,05M de redução de glosas tem payback lento. A adição do risco de pane no go-live (R$ 2,0M) empurra o projeto para o prejuízo (VPL negativo) em cenários de estresse de adoção realista ou pessimista.
* **Diretrizes para o Relatório Final (Aula 3):**
  • A recomendação pode ser NO-GO devido ao alto risco operacional ou GO Faseado (por exemplo, implantar primeiro apenas na UTI e faturamento). A discussão de governança deve pesar o risco de manter sistemas desintegrados vs. os custos da transição.""")
    ]
    
    for gid, title, guidelines in group_cases_details:
        add_heading_abnt(doc, title, level=2)
        
        # Add verbatim case study
        case_html = scenarios[gid]["caseStudy"]
        doc.add_paragraph().add_run("Transcrição Integral do Caso:").bold = True
        add_body_paragraph_with_html(doc, case_html, indent=True)
        
        # Add guidelines and parameters
        for line in guidelines.split("\n"):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.space_before = Pt(0)
            if line.strip().startswith("•") or line.strip().startswith("-"):
                # List item
                p.paragraph_format.left_indent = Cm(1.25)
                run = p.add_run(line.strip())
                apply_text_styling(run, font_name="Arial", size_pt=11.5)
            else:
                p.paragraph_format.first_line_indent = Cm(1.25)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run = p.add_run(line)
                apply_text_styling(run, font_name="Arial", size_pt=12)
        
    doc.add_page_break()

    # --- 4. BAREMA DE CORREÇÃO ---
    add_heading_abnt(doc, "4. Barema para Correção do Trabalho Final (10 Pontos)", level=1)
    add_body_paragraph(doc, "O professor utilizará os critérios abaixo para avaliar o Relatório Executivo final enviado pelas equipes:")
    
    # Rubric Table
    rubric_table = doc.add_table(rows=6, cols=5)
    rubric_table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    create_table_borders(rubric_table)
    
    headers_rubric = ["Critério", "Peso", "Excelente (2,0 pt)", "Regular (1,0 pt)", "Insuficiente (0 pt)"]
    col_widths_rubric = [Cm(2.5), Cm(1.2), Cm(4.1), Cm(4.1), Cm(4.1)]
    
    for idx, text in enumerate(headers_rubric):
        cell = rubric_table.rows[0].cells[idx]
        cell.width = col_widths_rubric[idx]
        p = cell.paragraphs[0]
        run = p.add_run(text)
        apply_text_styling(run, font_name="Arial", size_pt=9.5, bold=True, color_rgb=RGBColor(255, 255, 255))
        set_cell_background(cell, HEADER_BLUE_HEX)
        
    rubric_rows = [
        ("Conexão Contexto", "2,0 pt", "Extrai corretamente todos os dados e volumes de baseline do caso assistencial de forma precisa.", "Descreve o problema em linhas gerais, mas falha em quantificar de forma exata os prejuízos.", "Não contextualiza o cenário da instituição ou apresenta dados inventados não relacionados."),
        ("Modelagem de Risco", "2,0 pt", "Classifica rigorosamente os 4 tipos de riscos e propõe planos de mitigação práticos específicos.", "Lista os riscos de forma superficial ou propõe mitigantes genéricos sem detalhar fluxos de trabalho.", "Deixa de analisar uma ou mais categorias de riscos ou não propõe estratégias de contingência."),
        ("ROI sob Estresse", "2,0 pt", "Transcreve e analisa de forma exata as métricas financeiras nos 3 cenários, definindo o ponto de quebra.", "Preenche a tabela financeira, mas comete erros leves de transcrição ou falha em justificar os dados.", "A tabela apresenta dados em branco, erros grosseiros ou ausência de justificativa de estresse."),
        ("Recomendação", "2,0 pt", "Apresenta recomendação de GO/NO-GO fundamentada que cruza riscos clínicos com viabilidade financeira.", "Define a decisão, mas a justificativa é rasa, puramente intuitiva ou com fraco alinhamento aos dados.", "Ausência de decisão de recomendação ou falta absoluta de justificativa baseada em governança."),
        ("Qualidade de Escrita", "2,0 pt", "O relatório está integralmente preenchido, sem instruções do template e com linguagem clara de C-Level.", "Mantém trechos das instruções de preenchimento, erros gramaticais ou linguagem inadequada para diretoria.", "Relatório incompleto, desorganizado, sem o preenchimento de seções obrigatórias.")
    ]
    
    for row_idx, data in enumerate(rubric_rows):
        row = rubric_table.rows[row_idx + 1]
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            cell.width = col_widths_rubric[col_idx]
            p = cell.paragraphs[0]
            run = p.add_run(text)
            apply_text_styling(run, font_name="Arial", size_pt=9, color_rgb=SECONDARY_GRAY_RGB)
            if row_idx % 2 == 1:
                set_cell_background(cell, BG_LIGHT_HEX)
                
    doc.add_page_break()

    # --- 5. NOTAS DE FACILITAÇÃO SLIDE A SLIDE ---
    add_heading_abnt(doc, "5. Roteiro de Condução e Notas de Facilitador Slide a Slide", level=1)
    
    for c in classes:
        class_title = c["title"]
        add_heading_abnt(doc, class_title, level=2)
        for s in c["slides"]:
            slide_id = s["id"]
            slide_title = s["title"]
            raw_notes = s.get("notes", "Nenhuma nota de condução definida.")
            
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.left_indent = Cm(1.25)
            
            run_id = p.add_run(f"• {slide_id} ({slide_title}): ")
            apply_text_styling(run_id, font_name="Arial", size_pt=11.5, bold=True)
            
            # Clean and add notes text
            clean_note = raw_notes.replace("<strong>Fala sugerida:</strong>", "Fala sugerida:").replace("<strong>", "").replace("</strong>", "")
            run_note = p.add_run(clean_note)
            apply_text_styling(run_note, font_name="Arial", size_pt=11)
            
    # Save the template to the public directory
    output_dir = "/Users/faila/Library/CloudStorage/GoogleDrive-failapereira@gmail.com/My Drive/MOINHOS/public"
    os.makedirs(output_dir, exist_ok=True)
    file_path = os.path.join(output_dir, "Guia_Norteador_Professor.docx")
    doc.save(file_path)
    
    # Save a copy in the root workspace directory too
    root_path = "/Users/faila/Library/CloudStorage/GoogleDrive-failapereira@gmail.com/My Drive/MOINHOS/Guia_Norteador_Professor.docx"
    doc.save(root_path)
    
    print(f"Teacher Guide Word successfully generated at: {file_path}")

if __name__ == "__main__":
    main()
