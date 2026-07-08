import os
import docx
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Corporate Colors matching previous documents
BRAND_BLUE_RGB = RGBColor(2, 132, 199)
SECONDARY_GRAY_RGB = RGBColor(59, 81, 102)
MUTED_SLATE_RGB = RGBColor(107, 124, 147)

def set_cell_background(cell, color_hex):
    shd_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shd_xml))

def apply_text_styling(run, font_name="Arial", size_pt=11, bold=False, italic=False, color_rgb=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb
    else:
        run.font.color.rgb = RGBColor(0, 0, 0) # Default ABNT Black

def add_heading_academic(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.line_spacing = 1.15
    
    run = p.add_run(text)
    if level == 1:
        apply_text_styling(run, font_name="Arial", size_pt=12.5, bold=True, color_rgb=BRAND_BLUE_RGB)
    return p

def add_body_paragraph(doc, text="", bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    if text:
        run = p.add_run(text)
        apply_text_styling(run, font_name="Arial", size_pt=11, bold=bold, italic=italic)
    return p

def add_bullet_item(doc, bold_prefix, text, is_link=False):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    run_prefix = p.add_run(bold_prefix)
    apply_text_styling(run_prefix, font_name="Arial", size_pt=11, bold=True)
    
    if is_link:
        run_text = p.add_run(text)
        apply_text_styling(run_text, font_name="Arial", size_pt=11, color_rgb=BRAND_BLUE_RGB, bold=True)
    else:
        run_text = p.add_run(text)
        apply_text_styling(run_text, font_name="Arial", size_pt=11)
    return p

def main():
    doc = docx.Document()
    
    # Page setup - ABNT standard margins
    for section in doc.sections:
        section.top_margin = Cm(3.0)
        section.left_margin = Cm(3.0)
        section.bottom_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        
    # Header Logo / Institution block matching the screenshot
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_p.paragraph_format.space_before = Pt(0)
    logo_p.paragraph_format.space_after = Pt(12)
    if os.path.exists("Logo Moinhos.png"):
        logo_p.add_run().add_picture("Logo Moinhos.png", width=Inches(1.8))
        
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("MANUAL DE ACESSO AO AMBIENTE VIRTUAL")
    apply_text_styling(run_title, font_name="Arial", size_pt=14, bold=True, color_rgb=BRAND_BLUE_RGB)
    
    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(4)
    run_sub = sub_p.add_run("Diretrizes de Acompanhamento de Gravações de Aulas e Interações Individuais")
    apply_text_styling(run_sub, font_name="Arial", size_pt=11, italic=True, color_rgb=SECONDARY_GRAY_RGB)
    
    # Discipline
    disc_p = doc.add_paragraph()
    disc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disc_p.paragraph_format.space_after = Pt(4)
    run_disc = disc_p.add_run("Disciplina: Avaliação de Riscos e Investimentos com ROI")
    apply_text_styling(run_disc, font_name="Arial", size_pt=11, bold=True)
    
    # Teacher
    teach_p = doc.add_paragraph()
    teach_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    teach_p.paragraph_format.space_after = Pt(18)
    run_teach = teach_p.add_run("Profa. Dra. Faila Santos")
    apply_text_styling(run_teach, font_name="Arial", size_pt=11, bold=True)
    
    # Divider line
    div_p = doc.add_paragraph()
    div_p.paragraph_format.space_after = Pt(18)
    div_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div_run = div_p.add_run("―" * 45)
    div_run.font.color.rgb = SECONDARY_GRAY_RGB
    
    # Intro
    add_body_paragraph(doc, "Este manual destina-se aos discentes que necessitem realizar o acompanhamento acadêmico de forma assíncrona. A metodologia baseia-se em aprendizagem ativa síncrona. Portanto, a visualização da gravação deve ser rigorosamente integrada ao uso do portal virtual de simulação de ROI para que o aluno execute e valide os conceitos práticos apresentados.")
    
    # Section 1
    add_heading_academic(doc, "1. Acesso e Autenticação no Portal", 1)
    add_body_paragraph(doc, "Para acessar a interface do portal de aprendizagem e registrar suas interações individuais, execute o seguinte protocolo de autenticação:")
    add_bullet_item(doc, "Navegação Inicial: ", "Abra o navegador e acesse o endereço eletrônico do portal virtual ")
    # Add link paragraph style
    p_link = doc.paragraphs[-1]
    run_link = p_link.add_run("\nhttps://roi-learning-lab.onrender.com/")
    apply_text_styling(run_link, font_name="Arial", size_pt=11, bold=True, color_rgb=BRAND_BLUE_RGB)
    
    add_bullet_item(doc, "Seleção de Identidade: ", "Na tela de autenticação, clique no menu suspenso ('Selecione seu Nome') e selecione o seu nome correspondente.")
    add_bullet_item(doc, "Credencial de Acesso: ", "No campo 'Senha', insira o seu endereço de e-mail acadêmico cadastrado no MBA (ex: nome.sobrenome@hmv.org.br).")
    add_bullet_item(doc, "Autenticação: ", "Clique em 'Entrar' para liberar o acesso ao painel de navegação síncrona.")
    
    # Section 2
    add_heading_academic(doc, "2. Estudo Integrado com a Gravação", 1)
    add_body_paragraph(doc, "Recomenda-se posicionar a janela da gravação da videoaula e a janela do portal virtual lado a lado no seu monitor. Siga as orientações a seguir:")
    add_bullet_item(doc, "Navegação Concomitante: ", "Avance os slides teóricos no portal acompanhando a explanação da docente no vídeo.")
    add_bullet_item(doc, "Leituras Científicas e Referências: ", "Nos slides que contêm cards de artigos ou leituras recomendadas, utilize os links de apoio para consultar a literatura científica citada diretamente no portal.")
    
    # Section 3
    add_heading_academic(doc, "3. Realização e Entrega de Atividades Práticas (Trabalho de Equipe)", 1)
    add_body_paragraph(doc, "A avaliação das atividades de simulação ocorre de forma individual no portal. Mesmo estudando assincronamente, você deverá preencher o simulador referente à sua equipe:")
    add_bullet_item(doc, "Acesso ao Cenário: ", "Acesse a aba 'Trabalho em Grupo' para visualizar o estudo de caso clínico-operacional atribuído à sua equipe de estudos.")
    add_bullet_item(doc, "Modelagem Financeira e Riscos: ", "Preencha os campos da matriz de mitigação de riscos (Aula 1), os parâmetros de estresse do simulador de ROI (Aula 2) ou o parecer executivo do Business Case final (Aula 3).")
    add_bullet_item(doc, "Submissão: ", "Após completar o preenchimento, clique no botão 'Submeter' no rodapé do formulário.")

    # Save outputs to all relevant public directories
    output_dirs = [
        "/Users/faila/Library/CloudStorage/GoogleDrive-failapereira@gmail.com/My Drive/MOINHOS",
        "/Users/faila/Library/CloudStorage/GoogleDrive-failapereira@gmail.com/My Drive/MOINHOS/public",
        "/Users/faila/Library/CloudStorage/GoogleDrive-failapereira@gmail.com/My Drive/MOINHOS/deploy/public"
    ]
    
    filename = "Guia_Estudo_Assincrono_Estudante.docx"
    
    for o_dir in output_dirs:
        if os.path.exists(o_dir):
            path = os.path.join(o_dir, filename)
            doc.save(path)
            print(f"Document saved successfully: {path}")

if __name__ == "__main__":
    main()
