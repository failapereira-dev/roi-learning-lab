import os
import docx
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Corporate Colors matching previous documents
BRAND_BLUE_RGB = RGBColor(2, 132, 199)
SECONDARY_GRAY_RGB = RGBColor(59, 81, 102)
MUTED_SLATE_RGB = RGBColor(107, 124, 147)

def set_cell_background(cell, color_hex):
    shd_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shd_xml))

def apply_text_styling(run, font_name="Arial", size_pt=11, bold=False, italic=False, color_rgb=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb
    else:
        run.font.color.rgb = RGBColor(0, 0, 0)

def add_heading_academic(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.line_spacing = 1.15
    
    run = p.add_run(text)
    if level == 1:
        apply_text_styling(run, font_name="Arial", size_pt=12.5, bold=True, color_rgb=BRAND_BLUE_RGB)
    return p

def add_body_paragraph(doc, text="", bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    if text:
        run = p.add_run(text)
        apply_text_styling(run, font_name="Arial", size_pt=11, bold=bold, italic=italic)
    return p

def add_bullet_item(doc, bold_prefix, text, is_link=False):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    run_prefix = p.add_run(bold_prefix)
    apply_text_styling(run_prefix, font_name="Arial", size_pt=11, bold=True)
    
    if is_link:
        run_text = p.add_run(text)
        apply_text_styling(run_text, font_name="Arial", size_pt=11, color_rgb=BRAND_BLUE_RGB, bold=True)
    else:
        run_text = p.add_run(text)
        apply_text_styling(run_text, font_name="Arial", size_pt=11)
    return p

def main():
    doc = docx.Document()
    
    # Page setup - US Letter standard margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Header Logo
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_p.paragraph_format.space_before = Pt(0)
    logo_p.paragraph_format.space_after = Pt(12)
    if os.path.exists("Logo Moinhos.png"):
        logo_p.add_run().add_picture("Logo Moinhos.png", width=Inches(1.8))
        
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("STUDENT ACCESS MANUAL FOR THE VIRTUAL LAB")
    apply_text_styling(run_title, font_name="Arial", size_pt=14, bold=True, color_rgb=BRAND_BLUE_RGB)
    
    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(4)
    run_sub = sub_p.add_run("Guidelines for Asynchronous Participation, Recorded Sessions, and Individual Interactions")
    apply_text_styling(run_sub, font_name="Arial", size_pt=11, italic=True, color_rgb=SECONDARY_GRAY_RGB)
    
    # Course
    disc_p = doc.add_paragraph()
    disc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disc_p.paragraph_format.space_after = Pt(4)
    run_disc = disc_p.add_run("Course: Risk Assessment & Return on Investment (ROI) in Healthcare AI")
    apply_text_styling(run_disc, font_name="Arial", size_pt=11, bold=True)
    
    # Instructor
    teach_p = doc.add_paragraph()
    teach_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    teach_p.paragraph_format.space_after = Pt(18)
    run_teach = teach_p.add_run("Instructor: Faila Santos, PhD")
    apply_text_styling(run_teach, font_name="Arial", size_pt=11, bold=True)
    
    # Divider line
    div_p = doc.add_paragraph()
    div_p.paragraph_format.space_after = Pt(18)
    div_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div_run = div_p.add_run("―" * 45)
    div_run.font.color.rgb = SECONDARY_GRAY_RGB
    
    # Intro
    add_body_paragraph(doc, "This manual is designed for students who need to complete their coursework asynchronously. Our course is built on active learning methodologies. Therefore, you must integrate watching the recorded class sessions with hands-on practice in the virtual ROI portal to apply and validate the financial and operational concepts.")
    
    # Section 1
    add_heading_academic(doc, "1. Access and Authentication", 1)
    add_body_paragraph(doc, "To access the learning portal and record your interactions, execute the following authentication protocol:")
    add_bullet_item(doc, "Initial Navigation: ", "Open your browser and navigate to the virtual portal:")
    p_link = doc.paragraphs[-1]
    run_link = p_link.add_run("\nhttps://roi-learning-lab.onrender.com/") # fallback Render url
    apply_text_styling(run_link, font_name="Arial", size_pt=11, bold=True, color_rgb=BRAND_BLUE_RGB)
    
    add_bullet_item(doc, "Identity Selection: ", "On the login screen, click the dropdown menu ('Select your Name') and choose your corresponding student group name (e.g., Student Group 1).")
    add_bullet_item(doc, "Access Credential: ", "In the 'Password' field, enter your group name (e.g., student1 or student 1).")
    add_bullet_item(doc, "Login: ", "Click 'Login' to unlock access to the interactive class dashboard.")
    
    # Section 2
    add_heading_academic(doc, "2. Recorded Class Integration", 1)
    add_body_paragraph(doc, "We recommend positioning the video recording window and the virtual portal side-by-side on your monitor. Follow these guidelines:")
    add_bullet_item(doc, "Simultaneous Navigation: ", "Advance the theoretical slides in the portal while following the instructor's explanation in the video.")
    add_bullet_item(doc, "Scientific Literature & References: ", "For slides containing research links or recommended readings, click the links to review the scientific literature directly from the portal.")
    
    # Section 3
    add_heading_academic(doc, "3. Team Activities & Submissions", 1)
    add_body_paragraph(doc, "Evaluation of simulation activities is submitted individually on the portal. Even if studying asynchronously, you must fill out the simulator for your team:")
    add_bullet_item(doc, "Accessing the Scenario: ", "Click the 'Group Work' tab to review the clinical-operational case study assigned to your group.")
    add_bullet_item(doc, "Financial & Risk Modeling: ", "Fill in the risk mitigation matrix (Class 1), the stress parameters of the ROI simulator (Class 2), or the final Executive Business Case (Class 3).")
    add_bullet_item(doc, "Submission: ", "After completing all inputs, click the 'Submit' button at the bottom of the form to record your team's answers.")

    output_dir = "/Users/faila/Library/CloudStorage/GoogleDrive-failapereira@gmail.com/My Drive/MOINHOS/us/public"
    os.makedirs(output_dir, exist_ok=True)
    file_path = os.path.join(output_dir, "Student_Manual.docx")
    doc.save(file_path)
    print(f"US Student Manual successfully generated at: {file_path}")

if __name__ == "__main__":
    main()
